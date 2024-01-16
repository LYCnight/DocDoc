from LLMs import ChatGLM
from langchain.embeddings.huggingface import HuggingFaceEmbeddings

## 文档生成的方法
def DocGenerate(llm: ChatGLM,
                embedmodel: HuggingFaceEmbeddings
                ) -> str:
    '''
    传入目录和模型，返回文章全文（str 格式）
    '''
    # 0.config
    ProjectPath = "/remote-home/yy/lzd/DocDoc" 

    # 1.导入 LLM model 和 embedding model
    LLM = llm
    embeddings = embedmodel

    # test：
    print("---测试开始---")
    print(llm("你好"))
    print(embeddings.embed_query("北京大学"))
    print("---测试结束---")
    
    # 2.上传文件 loader
    # 提取word文档的文本
    from docx import Document
    filename = '汨罗江洪道治理报告书（2021）.docx'
    filepath = ProjectPath + "/test/testfile/" + filename
    doc = Document(filepath)   # load word
    # print(doc.paragraphs) # 打印内存地址
    texts = ""
    for paragraph in doc.paragraphs:
        # print(paragraph.text)  # 文本
        texts = texts + paragraph.text + "\n" # 拼接文本 
    # texts -> str
    print("step2 completed")

    # -----------------------------------------
    # 3.分割文档 spliter
    from langchain.text_splitter import SpacyTextSplitter
    from langchain.docstore.document import Document
    text_splitter = SpacyTextSplitter(
            separator = "\n\n",
            pipeline = "zh_core_web_sm",  # 加载 spacy 的中文模型
            max_length = 1_000_000,
            chunk_size=1000,    # 最大切割字符：1000个
            chunk_overlap = 200)
    
    texts_splited = text_splitter.split_text(texts)
    # texts_splited -> list[str]

    # 构造 langchain Document 对象
    text_Document_list = [Document(page_content=text, metadata={"from": "filename or book.txt"}) for text in texts_splited]
    # text_Document_list -> list[Document()]

    # ----------------------------------------
    # 4.向量化 embedding
    # 相似性搜索
    vs_path = "demo-vs"  # vector_db 存储地址
    from langchain.vectorstores import FAISS # FASIS
    # docs = embeddings.embed_documents(texts1) # -> list[向量]
    vector_store = FAISS.from_documents(text_Document_list, embeddings) # 建立 向量数据库
    vector_store.save_local(vs_path)  # 存储到vsdb中

    # ---------------------------------------
    # 5.向量检索 retriever
    vector_store = FAISS.load_local(vs_path, embeddings)  # 加载 vector_db
    related_docs_with_score = vector_store.similarity_search_with_score(query="在撰写《环评报告》的“结论”章节时，需要注意什么？", k=10) # 相似性搜索，搜索前K条知识并打分

    # ---------------------------------------
    # 6.模板拼接 prompt = content + query
    # 知识拼接
    contents = ""
    for pack in related_docs_with_score:
        doc, socre = pack
        content = doc.page_content
        print("检索到的知识=%s, from=%s, socre=%.3f"%(content, doc.metadata.get("from"), socre))
        contents += content  

    # 模板拼接
    query="帮我撰写的“结论”章节。结构要求：1. 项目概况 2. 区域环境质量现状评价结论 3. 环境影响分析结论 4. 环境可行性分析结论  5. 公众参与结论  6. 综合结论"
    template = f"已知：{contents}，请你根据你的知识和我提供给你的信息，{query}"

    # 7.生成回答 response
    response = LLM(template) # 基于知识的生成

    # 8.返回 response
    return response



def AIContinue(llm: ChatGLM,
               embedmodel: HuggingFaceEmbeddings, 
               contents: str = ""
               ) -> str:
    '''
    AI续写功能
    '''
    # 1.导入 LLM model 和 embedding model
    LLM = llm
    embeddings = embedmodel
    
    # 2.导入上下文 contents
    query="续写更多内容。（注意，请直接在原内容后续写新内容）"
    template = f"上下文：{contents}。 ---- 请你根据上下文，{query}"

    # 2.生成续写 response
    response = LLM(template) # 基于上下文的续写

    return response



def AIImprove( llm: ChatGLM,
               embedmodel: HuggingFaceEmbeddings, 
               contents: str = ""
               ) -> str:
    '''
    AI续写功能
    '''
    # 1.导入 LLM model 和 embedding model
    LLM = llm
    embeddings = embedmodel
    
    # 2.导入上下文 contents
    query="请你优化这段上下文，使其变得更像学术报告的风格"
    template = f"上下文：{contents}。 ---- {query}"

    # 2.生成续写 response
    response = LLM(template) # 基于上下文的续写

    return response