# 0.config
ProjectPath = "/remote-home/yy/lzd/DocDoc" 
MODEL_PATH = "/remote-home/share/LLM_model/chatglm3-6b"
TOKENIZER_PATH = "/remote-home/share/LLM_model/chatglm3-6b"
EMBEDDING_PATH = "/remote-home/share/LLM_model//bge-large-zh"


# 1.导入 LLM model 和 embedding model
from LLMs import ChatGLM
from typing import List

LLM = ChatGLM()
LLM.load_model(MODEL_PATH = MODEL_PATH, TOKENIZER_PATH = TOKENIZER_PATH)

from langchain.embeddings.huggingface import HuggingFaceEmbeddings
import numpy as np
embeddings = HuggingFaceEmbeddings(model_name = EMBEDDING_PATH,
                                    model_kwargs = {'device': "cuda"})

# 2.上传文件 loader
# 提取word文档的文本
from docx import Document
filename = '汨罗江洪道治理报告书（2021）.docx'
filepath = ProjectPath + "/test/testfile/" + filename
doc = Document(filepath)   # load word
# print(doc.paragraphs) # 打印内存地址
texts = ""
for paragraph in doc.paragraphs:
    # print(paragraph.text)  # 文本
    texts = texts + paragraph.text + "\n" # 拼接文本 
# texts -> str
print("step2 completed")

# -----------------------------------------
# 3.分割文档 spliter
from langchain.text_splitter import SpacyTextSplitter
from langchain.docstore.document import Document
text_splitter = SpacyTextSplitter(
        separator = "\n\n",
        pipeline = "zh_core_web_sm",  # 加载 spacy 的中文模型
        max_length = 1_000_000,
        chunk_size=1000,    # 最大切割字符：1000个
        chunk_overlap = 200)
 
texts_splited = text_splitter.split_text(texts)
# texts_splited -> list[str]

# 构造 langchain Document 对象
text_Document_list = [Document(page_content=text, metadata={"from": "filename or book.txt"}) for text in texts_splited]
# text_Document_list -> list[Document()]


# ----------------------------------------
# 4.向量化 embedding
# 相似性搜索
vs_path = "demo-vs"  # vector_db 存储地址
from langchain.vectorstores import FAISS # FASIS
# docs = embeddings.embed_documents(texts1) # -> list[向量]
vector_store = FAISS.from_documents(text_Document_list, embeddings) # 建立 向量数据库
vector_store.save_local(vs_path)  # 存储到vsdb中


# ---------------------------------------
# 5.向量检索 retriever
vector_store = FAISS.load_local(vs_path, embeddings)  # 加载 vector_db
related_docs_with_score = vector_store.similarity_search_with_score(query="在撰写《环评报告》的“结论”章节时，需要注意什么？", k=10) # 相似性搜索，搜索前K条知识并打分


# ---------------------------------------
# 6.模板拼接 prompt = content + query
# 知识拼接
contents = ""
for pack in related_docs_with_score:
    doc, socre = pack
    content = doc.page_content
    print("检索到的知识=%s, from=%s, socre=%.3f"%(content, doc.metadata.get("from"), socre))
    contents += content  

# 模板拼接
query="帮我撰写的“结论”章节。结构要求：1. 项目概况 2. 区域环境质量现状评价结论 3. 环境影响分析结论 4. 环境可行性分析结论  5. 公众参与结论  6. 综合结论"
template = f"已知：{contents}，请你根据你的知识和我提供给你的信息，{query}"


# 7.生成回答 response
response = LLM(template) # 基于知识的生成


# 8.生成 pdf/word/txt   - docgen
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Table, SimpleDocTemplate, Paragraph, Image  # 报告内容相关类
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfbase import pdfmetrics
from reportlab.lib import  colors
from reportlab.pdfbase.ttfonts import TTFont

pdfmetrics.registerFont(TTFont('宋体', './fonts/simsun.ttc')) # 字体注册

# 获取所有样式表
style = getSampleStyleSheet()
# 获取普通样式
ct = style['Normal']
ct.fontName = '宋体'

response_pdf = response.replace("\n", "<br/>") # 转换为 pdf 的换行符

# 生成pdf文件
pdfname = "Output8.pdf"
OUTPUT_PATH = "./output"
outputpath = OUTPUT_PATH + "/"+ pdfname 

doc = SimpleDocTemplate(outputpath, pagesize=letter)
doc.build([Paragraph(response_pdf, ct)])
print("*" * 100)
print(f"{pdfname}生成完毕！")
print(f"文件位于{outputpath}")


