#1 LLM model
from typing import Any, List
from transformers import AutoTokenizer, AutoModel

class chatGLM(): # ChatGLM 类
    def __init__(self, model_name, quantization_bit=4) -> None:
        self.tokenizer = AutoTokenizer.from_pretrained(model_name, trust_remote_code=True)
        model = AutoModel.from_pretrained(model_name, trust_remote_code=True).half().cuda().eval()
        self.model = model.quantize(quantization_bit)

    def __call__(self, prompt) -> Any:  # 生成回答的函数
        response, _ = self.model.chat(self.tokenizer , prompt) # 这里演示未使用流式接口. stream_chat()
        return response

llm =  chatGLM(model_name="THUDM/chatglm-6b") # 调用 __init__函数
prompt = "你好"
response = llm(prompt)  # 调用 __call__ 函数
print("response: %s"%response) # 你好👋！我是人工智能助手 ChatGLM-6B，很高兴见到你，欢迎问我任何问题。

# ----------------------------------------------------------
# 2 使用prompt模板，格式化生成新的prompt
from langchain import PromptTemplate
template = "请给我解释一下{query}的意思"
promptTem = PromptTemplate(input_variables=["query"], template=template)
prompt = promptTem.format(query="天道酬勤")

# ----------------------------------------------------------
# 3 用chain连链接llm和prompt组件
from langchain.chains.base import Chain
class DemoChain():
    def __init__(self, llm, prompt) -> None:
        self.llm = llm
        self.prompt = prompt

    def run(self, query, context=None) -> Any:
        if context is not None:  # context就是检索到的知识
            prompt = self.prompt.format(query=query, context=context) # 有知识则根据知识生成回答
        else:                   # 没有知识，则直接生成回答
            prompt = self.prompt.format(query=query)
        print("query=%s  -> prompt=%s"%(query, prompt))
        print("*"*60)
        response = self.llm(prompt) 
        return response
    
chain = DemoChain(llm=llm, prompt=promptTem)
print("-"*80)
chain.run(query="天道酬勤") # 没有知识的情况下生成回答
print("-"*80)

# ----------------------------------------------------------
# 4 示例 Embedding 和 vs-store
from langchain.embeddings.huggingface import HuggingFaceEmbeddings
import numpy as np
embeddings = HuggingFaceEmbeddings(model_name="GanymedeNil/text2vec-large-chinese",
                                    model_kwargs={'device': "cuda"})
query_result = embeddings.embed_query("天道酬勤") 
print("embedding query.shape=", np.array(query_result).shape)

# texts = """ '天道酬勤'并不是鼓励人们不劳而获，而是提醒人们要遵循自然规律，通过不断的努力和付出来追求自己的目标。\n这种努力不仅仅是指身体上的劳动，
# 也包括精神上的努力和思考，以及学习和适应变化的能力。\n只要一个人具备这些能力，他就有可能会获得成功。"""

# 提取word文档的文本
from docx import Document
path = '汨罗江洪道治理报告书（2021）.docx'
doc = Document(path)
# print(doc.paragraphs) # 打印内存地址
texts = ""
for paragraph in doc.paragraphs:
    # print(paragraph.text)  # 文本
    texts = texts + paragraph.text + "\n" # 拼接文本

# 文本分块
from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document
class TextSpliter(CharacterTextSplitter):  # 用于文本分块的类（按 '\n' 换行符进行切割）
    def __init__(self, separator: str = "\n\n", **kwargs: Any):
        super().__init__(separator, **kwargs)
    def split_text(self, text: str) -> List[str]:

        texts = text.split("\n")
        texts = [Document(page_content=text, metadata={"from": "filename or book.txt"}) for text in texts]
        return texts
text_splitter = TextSpliter()
texts = text_splitter.split_text(texts)
texts1 = [text.page_content for text in texts]

# 相似性搜索
vs_path = "demo-vs"  # vector_db 存储地址
from langchain.vectorstores import FAISS # FASIS
docs = embeddings.embed_documents(texts1) # embedddings: "GanymedeNil/text2vec-large-chinese"
vector_store = FAISS.from_documents(texts, embeddings)
vector_store.save_local(vs_path)  # 存储到vsdb中

vector_store = FAISS.load_local(vs_path, embeddings)  # 加载 vector_db
related_docs_with_score = vector_store.similarity_search_with_score(query="结论", k=10) # 相似性搜索，搜索前K条知识并打分

#-------------
# 5 基于查询到的知识做prompt 
context = ""
for pack in related_docs_with_score:
    doc, socre = pack
    content = doc.page_content
    print("检索到的知识=%s, from=%s, socre=%.3f"%(content, doc.metadata.get("from"), socre))
    context += content  # 知识拼接

# 重新配置一个基于上下文的模板再来调下语言模型
template = "已知{context}, 帮我写{query}的“结论”章节。结构要求：1. 项目概况 2. 区域环境质量现状评价结论 3. 环境影响分析结论 4. 环境可行性分析结论  5. 公众参与结论  6. 综合结论  "
promptTem = PromptTemplate(input_variables=["context", "query"], template=template)
chain = DemoChain(llm=llm, prompt=promptTem)
response = chain.run(query="《环评报告》", context=context) # 基于知识的生成
print("-"*80)
print(response)  
print("-"*80)


#--------------
# 6 将回答输出到pdf
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Table, SimpleDocTemplate, Paragraph, Image  # 报告内容相关类
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfbase import pdfmetrics
from reportlab.lib import  colors
from reportlab.pdfbase.ttfonts import TTFont

pdfmetrics.registerFont(TTFont('宋体', 'simsun.ttc')) # 字体注册

# 获取所有样式表
style = getSampleStyleSheet()
# 获取普通样式
ct = style['Normal']
ct.fontName = '宋体'

response_pdf = response.replace("\n", "<br/>") # 转换为 pdf 的换行符
# 生成pdf文件
pdfname = "report.pdf"
doc = SimpleDocTemplate(pdfname, pagesize=letter)
doc.build([Paragraph(response_pdf, ct)])
print(f"{pdfname}生成完毕！")



''' log
Loading checkpoint shards: 100%|██████████████████████████████████████████████████████████████████████████████████████████████████████████████████████████████████████████████████| 8/8 [00:16<00:00,  2.07s/it]
The dtype of attention mask (torch.int64) is not bool
response: 你好👋！我是人工智能助手 ChatGLM-6B，很高兴见到你，欢迎问我任何问题。
/opt/conda/envs/lzd/lib/python3.10/site-packages/langchain/__init__.py:34: UserWarning: Importing PromptTemplate from langchain root module is no longer supported. Please use langchain.prompts.PromptTemplate instead.
  warnings.warn(
--------------------------------------------------------------------------------
query=天道酬勤  -> prompt=请给我解释一下天道酬勤的意思
************************************************************
--------------------------------------------------------------------------------
No sentence-transformers model found with name GanymedeNil/text2vec-large-chinese. Creating a new one with MEAN pooling.
embedding query.shape= (1024,)
检索到的知识=采取上述措施后，本项目施工过程产生的固废均能得到合理处置，因而措施合理可行。, from=filename or book.txt, socre=893.371
检索到的知识=7.1 产业政策符合性分析     111108, from=filename or book.txt, socre=901.962
检索到的知识=②负责监督各子项目内所有施工期环保设施的日常运行管理，保障各环保设施的正常运行，并对环保设施的改进提出积极的建议。, from=filename or book.txt, socre=908.485
检索到的知识=本项目项目任务重、要求高、工期紧。为确保工程建设有序、有效地顺利进行，须充分利用以往类似工程的施工经验。工程开工前，编制严密的网络计划，严格按网络计划组织安排施工。编制计划充分考虑施工现场可能对工期造成延误的各种因素，确定进度作业指标时留有余地，一旦发生延误迅速采取补救措施。, from=filename or book.txt, socre=925.295
检索到的知识=b) 新增工程量：, from=filename or book.txt, socre=926.079
检索到的知识=③组织协调：配合当地环保部门，对环境工程建设质量、施工进度、投资的合理使用、环保设施运行等进行监督管理，确保各项措施落实到实处，发挥实效。, from=filename or book.txt, socre=926.211
检索到的知识=d)日常应该提高人员的素质，定期进行演习和必要的技术培训，掌握有关设备的性能，熟练使用，同时保证清污设备的技术先进和良好的工作状态。, from=filename or book.txt, socre=929.316
检索到的知识=③施工结束后进行生态恢复，采取的生态恢复景观效果应与景区自然人文景观相协调。, from=filename or book.txt, socre=940.933
检索到的知识=演习：应定期参加应急反应练习，检验应急计划中的各个环节是否能快速、协调有效的实施，提高应急反应系统的实战能力。演习目标：a) 使参与应急反应的各成员部分掌握、熟悉和深刻理解各自的职责；b) 保持应急反应个有关环节的快速、协调、有效的运作；c) 检查设备的可用性和性能；d) 考核各级应急反应人员对理论和实际操作技能的熟悉程度，并及时发现应急计划制定和实施过程中的问题和不足之处。, from=filename or book.txt, socre=942.652
检索到的知识=熟悉设计文件、了解现场地形变化情况，掌握当地的水文和相关的各项技术资料。组织项目部管理人员及工程技术人员，根据设计文件技术要求，施工现场的踏勘与调查结果，认真仔细地编制好专项施工方案,掌握本项目施工中的难点，与施工技术、施工工艺、明确安全施工，质量控制，文明施工规范制度，施工现场的质量、安全、文明施工、施工进度控制，工地环境卫生，并落实到每个施工人员。, from=filename or book.txt, socre=944.208
--------------------------------------------------------------------------------
query=《环评报告》  -> prompt=已知采取上述措施后，本项目施工过程产生的固废均能得到合理处置，因而措施合理可行。7.1 产业政策符合性分析    111108②负责监督各子项目内所有施工期环保设施的日常运行管理，保障各环保设施的正常运行，并对环保设施的改进提出积极的建议。本项目项目任务重、要求高、工期紧。为确保工程建设有序、有效地顺利进行，须充分利用以往类似工程的施工经验。工程开工前，编制严密的网络计划，严格按网络计划组织安排施工。编制计划充分考虑施工现场可能对工期造成延误的各种因素，确定进度作业指标时留有余地，一旦发生延误迅速采取补救措施。b) 新增工程量：③组织协调：配合当地环保部门，对环境工程建设质量、施工进度、投资的合理使用、环保设施运行等进行监督管理，确保各项措施落实到实处，发挥实效。d)日常应该提高人员的素质，定期进行演习和必要的技术培训，掌握有关设备的性能，熟练使用，同时保证清污设备的技术先进和良好的工作状态。③施工结束后进行生态恢复，采取的生态恢复景观效果应与景区自然人文景观相协调。演习：应定期参加应急反应练习，检验应急计划中的各个环节是否能快速、协调有效的实施，提高应急反应系统的实战能力。演习目标：a) 使参与应急反应的各成员部分掌握、熟悉和深刻理解各自的职责；b) 保持应急反应个有关环节的快速、协调、有效的运作；c) 检查设备的可用性和性能；d) 考核各级应急反应人员对理论和实际操作技能的熟悉程度，并及时发现应急计划制定和实施过程中的问题和不足之处。熟悉设计文件、了解现场地形变化情况，掌握当地的水文和相关的各项技术资料。组织项目部管理人员及工程技术人员，根据设计文件技术要求，施工现场的踏勘与调查结果，认真仔细地编制好专项施工方案,掌握本项目施工中的难点，与施工技术、施工工艺、明确安全施工，质量控制，文明施工规范制度，施工现场的质量、安全、文明施工、施工进度控制，工地环境卫生，并落实到每个施工人员。, 帮我写《环评报告》的“结论”章节。结构要求：1. 项目概况 2. 区域环境质量现状评价结论 3. 环境影响分析结论 4. 环境可行性分析结论  5. 公众参与结论  6. 综合结论  
************************************************************
尊敬的评审专家：

本报告旨在向您介绍我们项目的情况，包括项目概况、区域环境质量现状评价结论、环境影响分析结论、环境可行性分析结论、公众参与结论以及综合结论。

1. 项目概况

本项目是针对当地区域环境状况而进行改善的，旨在提高区域环境质量，减少污染排放，提高环境保护水平。项目的目标是建设一个环保设施完善、环境友好、资源利用合理的工程。

2. 区域环境质量现状评价结论

根据当前区域环境质量现状评价，当地区域的环境状况受到多种因素的影响，如工业污染、交通运输、农业和资源浪费等。这些环境问题的存在对当地人民群众的健康和生存造成了严重威胁，同时也对当地生态环境的平衡和稳定构成了威胁。

3. 环境影响分析结论

在施工过程中，我们将采取一系列环保措施，包括对施工过程进行监督、对环保设施进行日常运行管理、提高人员素质、加强生态恢复等。这些措施旨在减少对当地环境的影响，保护环境的可持续性。

4. 环境可行性分析结论

我们有信心，通过采取以上措施，本项目可以获得成功，不会对当地环境造成不可逆转的破坏。同时，我们也将积极采取措施，提高环境保护水平，减少污染排放，提高环境保护水平。

5. 公众参与结论

我们计划邀请当地公众参与项目，以确保项目的可持续性，减少对当地环境造成的影响。我们还将定期举办公众参与会议，向公众介绍项目的情况，听取公众的意见和反馈，及时解决公众提出的问题。

6. 综合结论

本报告综合了项目概况、区域环境质量现状评价结论、环境影响分析结论、环境可行性分析结论以及公众参与结论。我们有信心，通过以上措施，本工程可以实现可持续发展，保护环境的可持续性。
--------------------------------------------------------------------------------
'''